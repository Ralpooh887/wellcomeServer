from docx import Document
import os
import pypandoc

def create_word_document(data):
    """
    Webhook 데이터를 기반으로 Word 문서를 생성
    """
    document_name = f"record_{data['record_id']}.docx"
    document_path = os.path.join("/tmp", document_name)

    # Word 문서 생성
    doc = Document()
    doc.add_heading('Document for Record', level=1)
    doc.add_paragraph(f"Record ID: {data['record_id']}")
    doc.add_paragraph(f"Name: {data.get('name', 'N/A')}")
    doc.add_paragraph(f"Description: {data.get('description', 'N/A')}")
    doc.save(document_path)

    return document_path

def convert_to_pdf(word_file_path):
    """
    Word 파일을 PDF로 변환
    """
    pdf_file_path = word_file_path.replace('.docx', '.pdf')
    pypandoc.convert_file(word_file_path, 'pdf', outputfile=pdf_file_path)
    return pdf_file_path
